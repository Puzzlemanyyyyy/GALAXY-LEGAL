"""DOCX export for approved/active drafts.

Converts a (very light) subset of Markdown into a styled .docx document:
  - ``# Title`` -> Heading 1
  - ``## Section`` -> Heading 2
  - ``### Sub`` -> Heading 3
  - Bullet lines starting with ``- `` -> bullet list
  - ``**bold**`` and ``*italic*`` inline runs
  - ``[E:xxx]`` markers -> grey footnote-style runs

The file is uploaded to the private ``legal-documents`` bucket under
``<user_id>/<case_id>/exports/draft-<id>-v<ver>-<ts>.docx`` and a signed URL
is returned to the caller.
"""
from __future__ import annotations

import io
import re
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from services.supabase_client import get_supabase_admin


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITAL_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)")
_EV_RE   = re.compile(r"\[E:([A-Za-z0-9_-]+)\]")


def _render_paragraph(para, text: str, size: int = 11) -> None:
    """Render a single line of markdown-ish text into a docx paragraph.

    Handles bold/italic/evidence markers. Unknown formatting is written as
    plain text.
    """
    pos = 0
    tokens: list[tuple[str, str]] = []
    # Collect all matches; process left-to-right for the earliest match each step.
    while pos < len(text):
        candidates = []
        for kind, rx in (("bold", _BOLD_RE), ("ital", _ITAL_RE), ("ev", _EV_RE)):
            m = rx.search(text, pos)
            if m:
                candidates.append((m.start(), kind, m))
        if not candidates:
            tokens.append(("text", text[pos:]))
            break
        candidates.sort(key=lambda x: x[0])
        start, kind, m = candidates[0]
        if start > pos:
            tokens.append(("text", text[pos:start]))
        tokens.append((kind, m.group(1)))
        pos = m.end()

    for kind, body in tokens:
        if kind == "ev":
            # Render evidence markers as superscript so they read like a
            # footnote pointer, not as a low-contrast inline tag. Keep the
            # bracketed shape ("[E:e001]") so a copy-paste of the DOCX into
            # any other tool still surfaces the marker as plain text.
            run = para.add_run(f"[E:{body}]")
            run.font.size = Pt(max(7, size - 3))
            run.font.color.rgb = RGBColor(0x52, 0x59, 0x6A)  # ink-700, more legible than ink-600
            run.font.superscript = True
            continue
        run = para.add_run(body)
        run.font.size = Pt(size)
        if kind == "bold":
            run.bold = True
        elif kind == "ital":
            run.italic = True


def markdown_to_docx(title: str, content_md: str, *, footer_line: str | None = None) -> bytes:
    doc = Document()
    # Base style: Cormorant-like isn't installed system-wide; fall back to Calibri.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    for raw in content_md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _render_paragraph(p, line.lstrip()[2:])
        else:
            p = doc.add_paragraph()
            _render_paragraph(p, line)

    if footer_line:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        run = footer_para.add_run(footer_line)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0x9A, 0xA8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_draft_to_storage(*, draft: dict, owner_id: str) -> tuple[str, str]:
    """Render and upload a draft as DOCX. Returns (storage_path, signed_url)."""
    content_md = draft.get("content_md") or ""
    title = draft.get("title") or (draft.get("tipo_documento") or "Draft")
    version = draft.get("version") or 1
    case_id = draft["case_id"]
    draft_id = draft["id"]

    footer = f"Galaxy Legal · draft {draft_id} · v{version} · {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    blob = markdown_to_docx(title, content_md, footer_line=footer)

    path = f"{owner_id}/{case_id}/exports/draft-{draft_id}-v{version}-{uuid.uuid4().hex[:8]}.docx"
    admin = get_supabase_admin()
    admin.storage.from_("legal-documents").upload(
        path=path,
        file=blob,
        file_options={
            "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "upsert": "true",
        },
    )
    signed = admin.storage.from_("legal-documents").create_signed_url(path, 3600)
    url = signed.get("signedURL") or signed.get("signed_url") or signed.get("signedUrl")
    return path, url or ""
