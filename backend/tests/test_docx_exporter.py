"""Tests for the lightweight markdown → DOCX converter."""
import io

from docx import Document

from services.docx_exporter import markdown_to_docx


def test_docx_renders_headings_and_bullets():
    md = (
        "# Demanda civil — ACME\n"
        "## Hechos\n"
        "- Primero: El demandado incumplió el contrato [E:e001]\n"
        "- Segundo: **El importe** es de 1.234 EUR\n"
        "## Fundamentos\n"
        "Solo *texto corrido* aquí.\n"
    )
    blob = markdown_to_docx("Demanda civil — ACME", md, footer_line="Galaxy Legal · test")
    assert isinstance(blob, bytes) and len(blob) > 2000  # non-trivial docx
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs]
    # Title + headings + at least one bullet line
    assert any(t.startswith("Demanda civil") for t in texts)
    assert any("Hechos" == t for t in texts)
    assert any("Primero" in t and "incumpli" in t for t in texts)
    # Footer line made it through
    section = doc.sections[0]
    footer_text = "\n".join(p.text for p in section.footer.paragraphs)
    assert "Galaxy Legal" in footer_text


def test_docx_handles_empty_lines():
    blob = markdown_to_docx("Title", "# Title\n\nA paragraph.\n\n- one\n- two\n")
    doc = Document(io.BytesIO(blob))
    # Ensure no exception and at least the two bullets are rendered
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "one" in joined and "two" in joined


def test_docx_evidence_marker_is_superscript_with_brackets():
    """[E:xxx] markers must render as superscript and keep the bracketed shape.

    This is the contract with both lawyers (the marker should look like a
    footnote pointer, not a low-contrast inline tag) and with downstream
    tooling (any plain-text extraction of the docx must still surface the
    `[E:xxx]` pattern verbatim).
    """
    md = "Hechos: el demandado incumplió la cláusula [E:e001] el día 15.\n"
    blob = markdown_to_docx("Title", md)
    doc = Document(io.BytesIO(blob))

    # Plain-text extraction MUST still contain the bracketed marker verbatim.
    flat = "\n".join(p.text for p in doc.paragraphs)
    assert "[E:e001]" in flat, (
        "Evidence marker lost its bracketed shape on render. "
        "Tools doing plain-text grep on the .docx will break."
    )

    # Find the run that holds the marker and check its formatting.
    found = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == "[E:e001]":
                assert run.font.superscript is True, "marker run is not superscript"
                # Smaller than the body size (which defaults to 11pt → marker should be 7-8pt)
                assert run.font.size is not None and run.font.size.pt <= 8
                found = True
                break
        if found:
            break
    assert found, "Did not locate the evidence marker run in the document"
